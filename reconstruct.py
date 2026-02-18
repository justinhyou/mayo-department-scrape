import os
import docx
import pandas as pd
from collections import defaultdict
from docx.shared import RGBColor

input_csv = os.path.join(os.getcwd(), "input.csv")

possible_labels = ["Br", "Sp", "Peri", "Peds", "Vasc", "Other", "Skip"]


def screen(categories):
    """ Given the set of categories, find any unmatched labels. """
    unlabeled = set(categories)
    for label in possible_labels:
        if label in unlabeled:
            unlabeled.remove(label)

    return unlabeled


def regroup(df):
    """ From input csv, group entries by category. """
    regrouped = defaultdict(list)
    for i, row in df.iterrows():
        if str(row['title']) != "nan":
            recovered_entry = (row['title'], row['authors'], row['pub_info'], row['link'])
            category = row['category']
            regrouped[category].append(recovered_entry)
    return regrouped


def main():
    df = pd.read_csv(input_csv)
    regrouped = regroup(df)

    unlabeled = screen(list(regrouped.keys()))
    if unlabeled:
        print("Contains", unlabeled)
        return

    doc = docx.Document()

    for category in regrouped.keys():
        if category == "Skip":
            continue
        categorical = regrouped[category]
        p = doc.add_paragraph()
        run = p.add_run(category)
        run.bold = True

        for entry in categorical:
            p = doc.add_paragraph()

            run = p.add_run(entry[0] + " ")
            run.bold = False
            run.italic = False
            run.underline = True

            run = p.add_run(entry[1] + " ")
            run.bold = False
            run.italic = False

            run = p.add_run(entry[2] + " ")
            run.italic = True

            run = p.add_run(entry[3])
            run.underline = True
            run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

    doc.save("output.docx")


if __name__ == '__main__':
    main()
